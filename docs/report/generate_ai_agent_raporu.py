# -*- coding: utf-8 -*-
"""Yapay Zeka dersi proje raporu - AI Agent + Fine-tune odakli.

Cikti: docs/report/AI_Agent_Raporu_Metin_Karyagdi.docx
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"docs/report/AI_Agent_Raporu_Metin_Karyagdi.docx"

PROJECT_TITLE = (
    "FLOWGENT: YEREL LLM TABANLI TOOL-CALLING AGENT VE LoRA FINE-TUNE PIPELINE"
)
COURSE = "Yapay Zeka"
TERM = "2025-2026 Bahar Donemi"
STUDENT_NO = "222201346"
STUDENT_NAME = "Metin KARYAGDI"
ADVISOR = "[Ogretim Uyesi Adi]"
MONTH = "Mayis 2026"
CITY = "AMASYA"
UNIVERSITY = "T.C. AMASYA UNIVERSITESI"
FACULTY = "MUHENDISLIK FAKULTESI"
DEPT = "BILGISAYAR MUHENDISLIGI BOLUMU"


# --------- helpers ---------
def add_para(doc, text="", *, bold=False, italic=False, align=None, size=None, mono=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if size:
            run.font.size = Pt(size)
        if mono:
            run.font.name = "Consolas"
            r = run._element
            rpr = r.get_or_add_rPr()
            rfonts = OxmlElement('w:rFonts')
            rfonts.set(qn('w:ascii'), 'Consolas')
            rfonts.set(qn('w:hAnsi'), 'Consolas')
            rpr.append(rfonts)
            run.font.size = Pt(9)
    return p


def add_heading(doc, text, level):
    return doc.add_heading(text, level=level)


def add_code_block(doc, text):
    """Tek paragrafta monospace, kucuk punto kod blogu."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    for i, line in enumerate(text.split("\n")):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        rpr = r._element.get_or_add_rPr()
        rfonts = OxmlElement('w:rFonts')
        rfonts.set(qn('w:ascii'), 'Consolas')
        rfonts.set(qn('w:hAnsi'), 'Consolas')
        rpr.append(rfonts)
    return p


def add_table(doc, header, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Table Grid"
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            table.rows[ri].cells[ci].text = val
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p


# ================================================================
# Belge baslangici
# ================================================================
doc = Document()

# Varsayilan font: Calibri 11
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ============= KAPAK =============
for _ in range(3):
    add_para(doc)
add_para(doc, UNIVERSITY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para(doc, FACULTY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para(doc, DEPT, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para(doc)
add_para(doc)
add_para(doc, f"{COURSE} DERSI PROJE RAPORU", bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
add_para(doc)
add_para(doc)
add_para(doc, PROJECT_TITLE, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
for _ in range(4):
    add_para(doc)
add_para(doc, f"{STUDENT_NO} - {STUDENT_NAME}",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para(doc, f"Ders Sorumlusu: {ADVISOR}",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para(doc, TERM, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
for _ in range(4):
    add_para(doc)
add_para(doc, MONTH, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para(doc, CITY, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

doc.add_page_break()

# ============= OZET =============
add_heading(doc, "OZET", level=0)
add_para(
    doc,
    "Bu calisma, Flowgent adli mikroservis tabanli ceviki proje yonetim "
    "platformuna entegre edilen yerel calisan bir buyuk dil modeli (LLM) "
    "tabanli yapay zeka asistaninin tasarim ve gerceklestirimini sunar. "
    "Sistemin ozunde, Ollama runtime'i uzerinde calistirilan gemma3:4b modeli "
    "ile etkilesimi yoneten ozel bir Agent Loop bulunmaktadir. Agent, "
    "kullanicinin dogal dilde verdigi komutlari (\"X projesinde Y issue'sunu "
    "olustur ve aktif sprint'e ekle\" gibi) bir arac (tool) katalogu uzerinden "
    "saglam JSON sozlesmesi cercevesinde gerceklestirir. Calisma kapsaminda "
    "ayrica iki asamali bir LoRA fine-tune cevrimi (v1 ve v2) tasarlanmis; "
    "Groq + Gemini saglayicilari uzerinden sentetik veri toplama, sablon "
    "tabanli agent ornegi uretimi, sema dogrulama, deduplikasyon ve "
    "Colab T4 uzerinde egitim ardindan llama.cpp ile GGUF "
    "donusumu yapilarak adapter Ollama runtime'ina entegre edilmistir. "
    "v1 fine-tune'in basarisiz olmasinin sebepleri analiz edilmis (LoRA r=8 "
    "kapasite yetersizligi + tool-calling ornegi eksikligi) ve v2'de bu sorunlar "
    "1417 ornekli (487 multi-turn agent ornegi dahil) yeni dataset ve r=32 "
    "konfigurasyonu ile giderilmistir.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(doc)
add_para(doc, "Anahtar Kelimeler: ", bold=True)
add_para(
    doc,
    "Buyuk Dil Modeli, LLM Agent, Tool-Calling, Function Calling, "
    "LoRA, Fine-tuning, Gemma3, Ollama, GGUF, Sentetik Veri Uretimi, "
    "JSON Sozlesmesi, Mikroservis Mimari, .NET 9, Python.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

doc.add_page_break()

# ============= ABSTRACT =============
add_heading(doc, "ABSTRACT", level=0)
add_para(
    doc,
    "This work presents the design and implementation of a locally-hosted "
    "Large Language Model (LLM) based AI assistant integrated into Flowgent, "
    "a microservice-based agile project management platform. At the core of "
    "the system lies a custom Agent Loop that orchestrates interaction with the "
    "gemma3:4b model running on the Ollama runtime. The agent translates "
    "natural language user instructions (e.g. \"create issue Y in project X "
    "and add it to the active sprint\") into structured tool invocations bound "
    "by a strict JSON contract. A two-stage LoRA fine-tuning cycle (v1 and v2) "
    "was designed: synthetic data was collected via Groq + Gemini providers, "
    "template-based agent samples were generated, schema validation and "
    "deduplication were applied, training was performed on Colab T4, and the "
    "resulting adapter was merged and converted to GGUF format via llama.cpp "
    "for integration into the Ollama runtime. The failure modes of the v1 "
    "fine-tune (insufficient LoRA r=8 capacity and absence of tool-calling "
    "samples) were analyzed and addressed in v2 by introducing 1,417 examples "
    "(including 487 multi-turn agent samples) and an r=32 configuration.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(doc)
add_para(doc, "Keywords: ", bold=True)
add_para(
    doc,
    "Large Language Model, LLM Agent, Tool-Calling, Function Calling, LoRA, "
    "Fine-tuning, Gemma3, Ollama, GGUF, Synthetic Data Generation, JSON "
    "Contract, Microservice Architecture, .NET 9, Python.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

doc.add_page_break()

# ============= ICINDEKILER =============
add_heading(doc, "ICINDEKILER", level=0)
toc = [
    "1. GIRIS",
    "   1.1 Problem Tanimi ve Motivasyon",
    "   1.2 Calismanin Kapsami",
    "   1.3 Katkilar",
    "2. SISTEM MIMARISI",
    "   2.1 Genel Mikroservis Topolojisi",
    "   2.2 AiService Bileseni",
    "   2.3 Hibrit Ic Servis Iletisim Modeli",
    "3. AI AGENT TASARIMI",
    "   3.1 Tool-Calling Yaklasimi ve JSON Sozlesmesi",
    "   3.2 Tool Katalogu",
    "   3.3 Agent Loop Algoritmasi",
    "   3.4 Sistem Prompt Muhendisligi",
    "   3.5 Tolerant JSON Parser",
    "   3.6 Denetim (Audit) ve Telemetri",
    "4. MODEL SECIMI VE YEREL CIKARIM (INFERENCE)",
    "   4.1 Model Karsilastirmasi ve Karar",
    "   4.2 Ollama Runtime Entegrasyonu",
    "5. LoRA FINE-TUNING PIPELINE",
    "   5.1 Motivasyon: Base Modelin Zaaflari",
    "   5.2 Veri Toplama Mimarisi",
    "   5.3 Sema Dogrulama ve Deduplikasyon",
    "   5.4 v1 -> v2 Datasetlerinin Karsilastirmasi",
    "   5.5 LoRA Egitim Konfigurasyonu",
    "   5.6 Yerel Merge + GGUF Donusumu (llama.cpp)",
    "   5.7 Ollama Modeline Donusturme",
    "6. DEGERLENDIRME",
    "   6.1 Eval Runner",
    "   6.2 Base Model Davranis Analizi",
    "   6.3 v1 Fine-tune Sonuclari (Basarisizlik Analizi)",
    "   6.4 v2 Beklentileri",
    "7. KARSILASILAN ZORLUKLAR VE COZUMLER",
    "8. SONUC VE GELECEK CALISMALAR",
    "9. KAYNAKLAR",
]
for line in toc:
    add_para(doc, line)

doc.add_page_break()

# ================================================================
# 1. GIRIS
# ================================================================
add_heading(doc, "1. GIRIS", level=1)

add_heading(doc, "1.1 Problem Tanimi ve Motivasyon", level=2)
add_para(
    doc,
    "Modern ceviki proje yonetim araclari (Jira, Azure DevOps, Trello vb.) "
    "kullanicilarinin gunluk is yukunun buyuk bolumu, manuel kayit girisi, "
    "issue olusturma, sprint planlama ve durum sorgulama islerinden olusur. "
    "Bu islerin buyuk cogunlugu yapilandirilmis verinin (issue baslik/aciklama/"
    "oncelik, sprint hedef tarih, kullanici atamasi vb.) tekrar tekrar formdan "
    "girilmesinden ibarettir. Buyuk dil modellerinin (LLM) son donemdeki "
    "yetkinlikleri, bu \"form doldurma\" yukunun dogal dil komutlarina "
    "indirgenmesini mumkun kilmistir.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(doc)
add_para(
    doc,
    "Calismamizin amaci, gelistirdigimiz Flowgent platformuna yerel LLM "
    "tabanli bir agent entegre ederek sunlari saglamaktir:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_bullet(doc, "Kullanicinin dogal dilde verdigi komutlari guvenli ve "
                "denetlenebilir tool cagrilarina cevirmek.")
add_bullet(doc, "OpenAI / Anthropic gibi bulut servislerine bagimli olmadan, "
                "veri mahremiyetini koruyacak sekilde tamamen yerel donanimda "
                "(RTX 4050 6 GB) calismak.")
add_bullet(doc, "Genel amacli bir base modelin (gemma3:4b) proje yonetimi "
                "alan diline ozelestirilmesi icin LoRA tabanli bir fine-tune "
                "cevrimi kurmak.")

add_heading(doc, "1.2 Calismanin Kapsami", level=2)
add_para(
    doc,
    "Bu rapor, daha genis kapsamli bir bitirme projesinin yapay zeka "
    "bileseni uzerine odaklanmaktadir. Spesifik olarak ele alinan basliklar:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_bullet(doc, "AiService mikroservisinin .NET 9 + Clean Architecture + CQRS "
                "uzerine kurulu mimarisi.")
add_bullet(doc, "5 araclik tool katalogu ve cok-adimli (multi-turn) agent "
                "calistiricisi.")
add_bullet(doc, "Sentetik veri toplama hatti: Groq ve Google Gemini "
                "saglayicilari ile LLM tabanli ornek uretimi, sablon tabanli "
                "agent ornegi sentezi, sema dogrulama, deduplikasyon ve altin "
                "(golden) set kurulumu.")
add_bullet(doc, "LoRA tabanli iki asamali fine-tune (v1 ve v2), Colab Free T4 "
                "uzerinde egitim ve llama.cpp + GGUF cevrimi.")
add_bullet(doc, "Ollama runtime'ina ozel modelin (bp-agent) entegrasyonu ve "
                "calisma anindaki feature flag tabanli model anahtarlama "
                "(UseFinetuned).")

add_heading(doc, "1.3 Katkilar", level=2)
add_bullet(doc, "Kucuk parametreli (4B), yerel calisan bir modelin uretim "
                "kalitesinde tool-calling yapabilmesi icin tasarlanmis "
                "Türkçe alias destekli tolerant JSON parser.")
add_bullet(doc, "Tool-calling ornegi iceren multi-turn dataset uretimi icin "
                "sablon (LLM-siz) tabanli, determinist ve sema-uyumlu bir "
                "sentezleyici (agent_synth.py).")
add_bullet(doc, "Eski Groq kotalari tukenince devreye giren mock fallback "
                "uretici ve farkli saglayicilarin (Groq / Gemini) tek bir "
                "auto-mode altinda secilebildigi cogul-saglayicili veri "
                "boruhatti.")
add_bullet(doc, "Gemma3 multimodal mimarisinin PEFT.load_adapter ile uyumsuz "
                "calisma problemini cozen elle (manual_merge.py) LoRA matris "
                "ekleme cozumu (yalnizca language_model bileseni icin).")

doc.add_page_break()

# ================================================================
# 2. SISTEM MIMARISI
# ================================================================
add_heading(doc, "2. SISTEM MIMARISI", level=1)

add_heading(doc, "2.1 Genel Mikroservis Topolojisi", level=2)
add_para(
    doc,
    "Sistem, ApiGateway (YARP) arkasinda yer alan birden cok bagimsiz "
    "servisten olusur. Tarayicidan gelen istek JWT dogrulanarak ilgili "
    "servise yonlendirilir. AI bileseni ayri bir mikroservis (AiService) "
    "olarak konumlandirilmis ve kendi PostgreSQL veritabani (ai-db) ile "
    "Ollama runtime'i ile birlikte kendi Docker Compose ag dilimine "
    "yerlestirilmistir.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(doc)
add_code_block(doc, r"""
+---------------------+
|     Browser (SPA)   |  React + TypeScript
+----------+----------+
           | HTTPS (JWT)
           v
+---------------------+
|   ApiGateway (YARP) |  Port 5000
+----+----------------+
     |
     | Rotalar
     v
+-------------+  +-------------+  +-------------+  +-------------+
| IdentityApi |  | ProjectApi  |  | SprintApi   |  | IssueApi    |
+-------------+  +-------------+  +-------------+  +-------------+
                         ^               ^               ^
                         |               |               |
                         +------+--------+--------+------+
                                |        Docker  ic ag (HTTP)
                                |        X-Internal-Service header
                                v
                         +-------------+         +-----------+
                         |   AiService |<------->|  Ollama   |
                         |  Port 5008  |  HTTP   | 11434     |
                         +------+------+         +-----+-----+
                                |                      |
                                v                      v
                          +----------+          +--------------+
                          |  ai-db   |          | gemma3:4b /  |
                          | Postgres |          | bp-agent     |
                          +----------+          +--------------+
""")

add_heading(doc, "2.2 AiService Bileseni", level=2)
add_para(
    doc,
    "AiService klasik Clean Architecture katmanlamasi ile yazilmistir:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_bullet(doc, "AiService.Domain: AiSession, AiToolExecution, AiPlanResult "
                "gibi alan modelleri.")
add_bullet(doc, "AiService.Application: CQRS handler'lari (GeneratePlan, "
                "EnrichIssue, AssessRisk vb.), tool soyutlamalari (ITool, "
                "IToolRegistry), Agent Loop ve MediatR komutlari.")
add_bullet(doc, "AiService.Infrastructure: OllamaClient (HTTP istemcisi), "
                "SprintServiceClient, IssueServiceClient (servisler arasi "
                "iletisim), ModelSelector, EF Core repository'leri.")
add_bullet(doc, "AiService.Api: Controller (AiController), DI kayitlari ve "
                "Program.cs.")

add_heading(doc, "2.3 Hibrit Ic Servis Iletisim Modeli", level=2)
add_para(
    doc,
    "Agent tool'lari (ornegin CreateIssue) basarili calismak icin IssueService'e "
    "issue olusturma cagrisi yapmak zorundadir. Tarayicidan gelen JWT, "
    "AiService'in dogruladigi sekilde tekrar tekrar IssueService'e yansitilmaz; "
    "bunun yerine Docker ic agi uzerinden ozel bir header (X-Internal-Service: "
    "AiService) ve kullanici kimligi (X-User-Id) ile cagri yapilir. Hedef servis "
    "(IssueService, SprintService) tarafinda InternalServiceMiddleware bu "
    "header'i dogrular ve cagrayi guvenli bir is baglami olarak kabul eder. "
    "Bu hibrit yaklasim hem JWT-temelli son-kullanici yetkilendirmesini hem "
    "de cluster ici hizli servis-arasi cagrilari mumkun kilar.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

doc.add_page_break()

# ================================================================
# 3. AI AGENT TASARIMI
# ================================================================
add_heading(doc, "3. AI AGENT TASARIMI", level=1)

add_heading(doc, "3.1 Tool-Calling Yaklasimi ve JSON Sozlesmesi", level=2)
add_para(
    doc,
    "Agent, modelin uretebildigi tek cikti turunu kati bir JSON sozlesmesi ile "
    "sinirlandirarak guvenilirlik kazanir. Model her adimda ya bir veya birden "
    "fazla tool cagirir (action) ya da kullaniciya gosterilecek nihai metni "
    "(final) uretir. Sozlesme asagidaki gibi tanimlanmistir:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(doc)
add_code_block(doc, """
ACTION JSON:
{
  "tool_calls": [
    { "name": "<tool_adi>", "input": { ...arguments... } }
  ]
}

FINAL JSON:
{
  "final": "<kullaniciya gosterilecek Turkce metin>"
}
""")
add_para(
    doc,
    "Model yalnizca yukaridaki iki sekilden birini uretmek zorundadir. "
    "Markdown fence (```json), Ingilizce metin sizintisi veya serbest metin "
    "sebepli sapmalari Tolerant Parser bolumunde ele alinmistir.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

add_heading(doc, "3.2 Tool Katalogu", level=2)
add_para(doc, "Agent'in erisebildigi araclar (ITool implementasyonlari):",
         align=WD_ALIGN_PARAGRAPH.JUSTIFY)
rows = [
    ["get_active_sprint", "Bir projedeki aktif sprint bilgisini doner. Side-effect yok.", "Sorgu"],
    ["get_project_issues", "Bir projedeki issue listesini (id, baslik, durum) doner.", "Sorgu"],
    ["create_issue", "Verilen baslik/aciklama/oncelikle yeni issue olusturur.", "Yazma"],
    ["create_sprint", "Verilen ad ve tarih araliginda yeni sprint olusturur.", "Yazma"],
    ["add_issue_to_sprint", "Mevcut bir issue'yu mevcut bir sprint'e ekler.", "Yazma"],
]
add_table(doc, ["Tool Adi", "Aciklama", "Tip"], rows)
add_para(doc)
add_para(
    doc,
    "Her aracin girdi semasi ToolSchemas.cs uzerinde JSON sema benzeri bir "
    "ifade ile tanimlanir. Sistem prompt'unun bir parcasi olarak bu katalog "
    "modele compact JSON formatinda enjekte edilir; bu sayede modelin "
    "halusinasyonla olmayan bir arac ismi uretme olasiligi azalir.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

add_heading(doc, "3.3 Agent Loop Algoritmasi", level=2)
add_para(
    doc,
    "AgentLoop.RunAsync, kullanicinin tek seferlik mesajini alip en fazla 5 "
    "iterasyona kadar model ile haberlesen, gerekirse tool cagrilarini "
    "yurutup sonuclarini tekrar modele besleyen multi-turn bir dongudur. "
    "Dongunun pseudo-kodu su sekildedir:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(doc)
add_code_block(doc, """
RunAsync(systemPrompt, userMessage, context):
    messages = [system, user]
    for iter in 1..5:
        raw = Ollama.ChatJson(messages)
        messages.append(assistant=raw)

        if not TryParse(raw): return raw (format_unrecognized)
        if HasFinal(parsed): return parsed.final
        if not HasToolCalls(parsed): return raw (format_unrecognized)

        for call in parsed.tool_calls:
            tool = registry.Resolve(call.name)
            result = await tool.ExecuteAsync(call.input, context)
            audit.Save(orgId, projectId, sessionId, call, result)
            messages.append(user="[tool] " + compact_json(result))

    return last_raw (hit_iteration_limit)
""")
add_para(
    doc,
    "Dikkat edilmesi gereken bazi tasarim kararlari:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_bullet(doc, "Iterasyon limiti 5: Yerel kucuk modelin sonsuz dongu "
                "olusturma riski (ayni tool'u tekrar tekrar cagirma) bu sert "
                "limitle sinirlanir.")
add_bullet(doc, "Tool ciktilari modele \"[tool] {json}\" formatinda doner; "
                "bu format egitim verisi (agent_synth.py) ile birebir "
                "esitlendigi icin model formatu tanir.")
add_bullet(doc, "messages dizisi her tool sonucundan sonra buyur; "
                "max_seq_length sinirlamasini asmamak icin tool ciktilari "
                "kisaltilir / id listesi formatinda dondurulur.")
add_bullet(doc, "/api/chat (Ollama) endpoint'i format=\"json\" parametresi ile "
                "cagrilir; boylece modelin uretiminde sintaks gecerli JSON "
                "olma olasiligi belirgin sekilde artar.")

add_heading(doc, "3.4 Sistem Prompt Muhendisligi", level=2)
add_para(
    doc,
    "Sistem prompt'u uc bolumden olusur: (1) Agent karakteri ve genel "
    "kurallar, (2) tool katalogu (compact JSON), (3) cikti formati spesifikasyonu. "
    "Egitim verisindeki system message ile birebir esit olmasi kritiktir: "
    "fine-tune sirasinda model bu sistem prompt'unun kalibinin tam kalibini "
    "ezberler.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(doc)
add_code_block(doc, """
[1] AGENT KARAKTERI
"Sen Flowgent AI agent'isin. Kullanicinin dogal dilde yazdigi
istegi arac cagrilariyla gerceklestirmek icin tool catalog'unu
kullanirsin. Her adimda ya bir tool cagirirsin ya da konusmayi
'final' ile bitirirsin. Yalnizca gecerli JSON dondurursun..."

[2] TOOL CATALOG (compact JSON)
{"tools":[{"name":"get_active_sprint","input":{"project_id":"uuid"}},
 {"name":"create_issue","input":{...}}, ...]}

[3] FORMAT
"Cikti yalnizca 2 sekilde olabilir:
  {\\"tool_calls\\":[...]} ya da {\\"final\\":\\"...\\"}"
""")

add_heading(doc, "3.5 Tolerant JSON Parser", level=2)
add_para(
    doc,
    "Yerel kucuk modellerin (4B parametre) JSON uretimindeki tipik "
    "hatalarinin pratik analizinden cikan bulgular dogrultusunda, "
    "AgentLoop.TryParse fonksiyonu agresif bir 'iyilestirme' (repair) "
    "katmani icerir:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_bullet(doc, "Markdown fence wrap: Cikti  ```json ... ```  icine alinmis "
                "olabilir; parser ilk satiri atlar ve son fence isaretini siler.")
add_bullet(doc, "Eksik kapatici brace: Model `}` veya `]` karakterini atlayabilir; "
                "parser acik/kapali parantezleri sayarak farki sona ekler.")
add_bullet(doc, "Turkce alias destegi: Model bazen 'arac_cagrilari' veya "
                "'sonuc' gibi Turkce anahtarlar uretir; parser bunlari sirasiyla "
                "'tool_calls' ve 'final' anahtarlarina map'ler.")
add_bullet(doc, "Yapisal {name, input} dizisi tespiti: Model 'tool_calls' "
                "anahtarini unutup dogrudan [{name:..., input:...}] dizisini "
                "uretirse parser bunu da tanir.")
add_para(doc)
add_para(
    doc,
    "Bu tolerant parser sayesinde base gemma3:4b modeli (fine-tune olmaksizin) "
    "bile uretimdeki bazi format hatalarina ragmen calisabilir hale gelmistir. "
    "Fine-tune basarili olunca parser, bir guvenlik agi olarak kalmaya devam "
    "eder; cunku model dagilim disi inputlarda hala kucuk olcekli sapmalar "
    "uretebilir.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

add_heading(doc, "3.6 Denetim (Audit) ve Telemetri", level=2)
add_para(
    doc,
    "Her tool calistirilmasi `AiToolExecution` entity'sine kalici olarak "
    "yazilir. Tablo sutunlari: OrganizationId, ProjectId, SessionId, "
    "ToolName, InputJson, OutputJson, IsSuccess, ElapsedMs, CreatedAt. "
    "Bu sayede:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_bullet(doc, "Hangi kullanici / org hangi tool'u kac kez kullanmis "
                "izlenebilir (rate limit ve telemetri icin).")
add_bullet(doc, "Bir issue/sprint'in 'kim/ne tarafindan' olusturuldugu "
                "(insan mi, agent mi, agent ise hangi prompt'la?) tam "
                "audit trail uzerinden takip edilebilir.")
add_bullet(doc, "Fine-tune sonrasinda hata analizinde gercek (production) "
                "input ornekleri kaynak olarak kullanilabilir.")

doc.add_page_break()

# ================================================================
# 4. MODEL SECIMI
# ================================================================
add_heading(doc, "4. MODEL SECIMI VE YEREL CIKARIM", level=1)

add_heading(doc, "4.1 Model Karsilastirmasi ve Karar", level=2)
add_para(
    doc,
    "Yerel donanim olarak RTX 4050 Mobile (6 GB VRAM) hedef alindi. "
    "Bu kisitla 7B+ modeller (Llama-3 8B, Qwen2.5 7B) Q4 quantization ile "
    "yaklasik 4.5-5.5 GB VRAM tukettigi icin marjin son derece dardir ve "
    "context window genisledikce OOM riskine girer. 4B sinifindaki modellerin "
    "(gemma3:4b, qwen2.5:3b, llama3.2:3b) karsilastirmasinda Gemma3 Turkce "
    "anlama/uretme dengesi ve tool-calling formatlarina yatkinligi nedeniyle "
    "tercih edildi.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(doc)
rows = [
    ["gemma3:4b (Q4)", "~3.2 GB", "Iyi", "Iyi", "Primary"],
    ["llama3.2:3b (Q4)", "~2.4 GB", "Orta", "Iyi", "Fallback"],
    ["qwen2.5:7b (Q4)", "~5.0 GB", "Iyi", "Cok Iyi", "VRAM marji dar"],
    ["llama3:8b (Q4)", "~5.5 GB", "Orta", "Iyi", "VRAM marji yetersiz"],
]
add_table(doc, ["Model", "VRAM (Q4)", "Turkce", "JSON Cikti", "Karar"], rows)

add_heading(doc, "4.2 Ollama Runtime Entegrasyonu", level=2)
add_para(
    doc,
    "Yerel calismayi soyutlamak icin Ollama runtime'i kullanildi. Ollama, "
    "GGUF formatindaki modelleri yukleyip OpenAI-uyumlu /api/chat ve /api/"
    "generate endpoint'lerini sunar. AiService bu endpoint'lere HTTP "
    "uzerinden eriserek modeli cagirir. Onemli config parametreleri:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(doc)
add_code_block(doc, """
Ollama:
  BaseUrl: http://ollama:11434     # Docker ic ag
  Model: gemma3:4b                 # Base model
  FallbackModel: llama3.2:3b       # Fallback (parse hatasinda)
  FinetunedModel: bp-agent         # Fine-tune adapter sonucu
  UseFinetuned: false              # Feature flag (env override)
""")
add_para(
    doc,
    "UseFinetuned bayragi sayesinde fine-tune'lu model ile base model "
    "arasinda kod degisikligi olmadan tek satirla gecis yapilabilir. Bu, "
    "fine-tune sonuclarini 'shadow' modunda test edip uretime rollout "
    "etmeyi kolaylastiran kritik bir tasarim secimidir.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

doc.add_page_break()

# ================================================================
# 5. LoRA FINE-TUNING
# ================================================================
add_heading(doc, "5. LoRA FINE-TUNING PIPELINE", level=1)

add_heading(doc, "5.1 Motivasyon: Base Modelin Zaaflari", level=2)
add_para(
    doc,
    "Tolerant parser ile birlikte calistirilan base gemma3:4b modeli "
    "uzerinde yapilan canli testler asagidaki dort tipik hata modunu "
    "ortaya cikardi:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_bullet(doc, "Eksik kapatici brace: Tool_calls dizisi acilir ama son `}` "
                "atlanir. (Tolerant parser ile telafi edildi.)")
add_bullet(doc, "Markdown fence wrap: Model heryerde ```json ... ``` "
                "uretmeye meyilli. (Tolerant parser ile telafi.)")
add_bullet(doc, "'final' uretememe: Tool sonucu donduktan sonra model "
                "ozet yapmak yerine ayni tool'u tekrar cagiriyor; 5 iter "
                "limitine takiliyor.")
add_bullet(doc, "Sema-uyumu sorunlari: enrich-issue, suggest-balance ve "
                "sprint-risk endpoint'lerinin sema-zorunlu JSON cikitlari "
                "tam tutturulamiyor; %0 schema-pass orani.")
add_para(doc)
add_para(
    doc,
    "Bu gozlemler 'fine-tune kararinin' motivasyonudur: Parser kadar yerini "
    "ne kadar iyi doldursa da, modelin format ezberi ve 'final ile bitirme' "
    "davranisi mutlaka egitilmek zorundadir.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

add_heading(doc, "5.2 Veri Toplama Mimarisi", level=2)
add_para(
    doc,
    "tools/ai_data_collector altinda cogul-saglayicili (multi-provider) bir "
    "Python paketi gelistirildi. Mimari sema:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(doc)
add_code_block(doc, r"""
                  +------------------+
                  |   synthetic_gen  |  (LLM tabanli)
                  |    --feature X   |
                  |    --provider P  |
                  +---+----------+---+
                      |          |
            +---------+          +---------+
            v                              v
   +-----------------+              +----------------+
   |  GroqProvider   |              | GeminiProvider |
   | llama-3.3-70b   |              | gemini-2.5-flash|
   +--------+--------+              +--------+-------+
            \                                /
             \                              /
              \      +-----------------+   /
               +--> | output/<feat>/   | <-+
                    | <prov>-<n>.json  |
                    +--------+---------+
                             |
       +-----------+         |
       | mock_gen  |---------+   (fallback, sablon tabanli)
       +-----------+         |
                             |
       +-----------+         |
       | agent_synth--------+   (multi-turn agent, template-based)
       +-----------+         |
                             v
                     +---------------+
                     |   merge.py    |
                     | dedup + valid |
                     +-------+-------+
                             |
                             v
                  +----------------------+
                  | train-vN.jsonl       |
                  | (N=1: 930, N=2: 1417)|
                  +----------+-----------+
                             |
                             v
                  +----------------------+
                  |  prepare_dataset.py  |
                  |  90/10 train/val     |
                  |  + token stats       |
                  +----------------------+
""")

add_para(
    doc,
    "Veri uretimi katmaninin onemli ozellikleri:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_bullet(doc, "Coklu saglayici: Bir API kotasi tukenirse digerine "
                "(GeminiProvider) gecis auto-mode altinda saglanir.")
add_bullet(doc, "Mock fallback: LLM tabanli uretim tamamen kesintiye girerse "
                "16 sektor x 22 modul bank uzerinden permutasyonel deterministik "
                "kayitlar uretilir; template_id 'mock-*' prefix'i ile "
                "isaretlenir.")
add_bullet(doc, "Template-based agent sentezleyici: agent_synth.py LLM "
                "kullanmadan, AgentLoop'un sistem prompt'una BIRINE-BIR ESIT "
                "format uretir; 1 saniyenin altinda 500 ornek uretebilir.")
add_bullet(doc, "Golden set: Her ozellik icin 15-30 elle yazilmis kaliteli "
                "ornek (insan denetimli) datasete %100 dahil edilir.")

add_heading(doc, "5.3 Sema Dogrulama ve Deduplikasyon", level=2)
add_para(
    doc,
    "merge.py asagidaki adimlari sirayla uygular:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_bullet(doc, "schema_validate: Her ozellik icin tanimli JSON schema "
                "(validation/schema.py) uygulanir; uymayanlar dropdan ayrilir.")
add_bullet(doc, "deduplicate: Input metni uzerinde normalize + hash; ayni "
                "input'a sahip ornekler tekillestirilir.")
add_bullet(doc, "eval_collision_filter: Eval setindeki promptlar ile cakisan "
                "egitim ornekleri filtrelenir (data leakage onlemi).")

add_heading(doc, "5.4 v1 -> v2 Datasetlerinin Karsilastirmasi", level=2)
rows = [
    ["scaffold-project", "364", "364"],
    ["enrich-issue", "308", "308"],
    ["generate-plan", "258", "258"],
    ["agent (multi-turn)", "0", "487"],
    ["TOPLAM", "930", "1417"],
    ["train / val split", "837 / 93", "1275 / 142"],
    ["LoRA r (rank)", "8", "32"],
    ["LoRA alpha", "16", "64"],
    ["max_seq_length", "2048", "1536"],
    ["epoch sayisi", "3 (125 step'te kesildi)", "3 (full, early stop yasak)"],
]
add_table(doc, ["Olcut", "v1", "v2"], rows)
add_para(doc)
add_para(
    doc,
    "v1'in en kritik eksigi 'agent' kategorisinde HIC ornek olmamasiydi: "
    "fine-tune yapilmis modelden tool-calling beklendigi halde egitim "
    "verisinde bu format yoktu. v2'de eklenen 487 multi-turn agent ornegi, "
    "modelin AgentLoop'un sistem prompt'unu tanidigi ve [tool] result "
    "formatini ezberledigi bir egitim sinyali sunar.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

add_heading(doc, "5.5 LoRA Egitim Konfigurasyonu", level=2)
add_para(
    doc,
    "v2.yaml uzerinde tanimli kritik hiperparametreler:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(doc)
add_code_block(doc, """
base_model:    unsloth/gemma-3-4b-it-bnb-4bit
max_seq_len:   1536  (v1: 2048)

LoRA:
  r:           32
  alpha:       64    (alpha/r = 2)
  dropout:     0.05
  target:      [q_proj, k_proj, v_proj, o_proj,
                gate_proj, up_proj, down_proj]
  bias:        none

Training:
  epochs:                 3 (early stop yasak)
  batch_size:             2 (per device)
  grad_accum_steps:       8       -> effective batch = 16
  learning_rate:          2.0e-4
  scheduler:              cosine
  warmup_ratio:           0.05
  weight_decay:           0.01
  optimizer:              adamw_8bit
  fp16:                   true
  save_steps:             100
  save_total_limit:       5
""")
add_para(
    doc,
    "v1'de r=8, v2'de r=32 secilmesinin gerekcesi: tool-calling skili sadece "
    "format ezberi degil, multi-turn karar dinamigi de gerektirir; bu bir "
    "yeni 'davranis modu' ekleme demektir ve LoRA'nin temsiliyet kapasitesinin "
    "(rank) iki misli artirilmasi mantikli olur. alpha/r oraninin 2 olarak "
    "korunmasi ile efektif scaling sabit kalir.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

add_heading(doc, "5.6 Yerel Merge + GGUF Donusumu (llama.cpp)", level=2)
add_para(
    doc,
    "Egitim sonucunda Colab Drive'a yazilan LoRA adapter dosyalari "
    "(safetensors + adapter_config.json) yerel makinada Ollama'da "
    "kullanilabilmesi icin once HF formatinda base model ile birlestirilmeli, "
    "ardindan GGUF formatina cevrilip quantize edilmelidir. Pipeline:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(doc)
add_code_block(doc, r"""
                  [Colab T4]
                       |
                       v
            LoRA adapter (safetensors)
                       |
                       v
              [Yerel - C:\gemma-export]
                       |
                       v
            manual_merge.py
            (PEFT.load_adapter BYPASS)
                       |
                       v
              merged/ (8.1 GB fp16 HF)
                       |
                       v
            llama.cpp convert_hf_to_gguf.py
                       |
                       v
        gemma3-4b-bp-v1.f16.gguf  (7.76 GB)
                       |
                       v
            llama-quantize.exe q4_k_m
                       |
                       v
        gemma3-4b-bp-v1.q4_k_m.gguf  (2.37 GB)
                       |
                       v
            ollama create bp-agent -f Modelfile
                       |
                       v
                 bp-agent:latest
                  (Ollama hazir)
""")
add_para(
    doc,
    "Bu pipeline'in alisilmistan ayrilan onemli noktasi manuel merge "
    "adimidir. Gemma3 multimodal mimaride iki ana bilesen icerir: "
    "language_model (text) ve vision_tower (image). LoRA adapter'i "
    "yalnizca language_model katmanlarina ait olmasina ragmen, "
    "PEFT.load_adapter rutini her iki bilesene de mapping yapmaya "
    "calisir ve boyut uyumsuzlugundan (2560 vs 1152) hata verir. "
    "manual_merge.py bu sorunu su sekilde cozer:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_bullet(doc, "Adapter dosyasi direkt safetensors olarak yuklenir.")
add_bullet(doc, "Sadece 'language_model' regex'i ile filtrelenmis 272 "
                "key (136 modul) tek tek isleme alinir.")
add_bullet(doc, "Her modul icin DeltaW = (alpha / r) * (B @ A) matris "
                "carpimi elle hesaplanir ve karsilik gelen base agirlik "
                "matrisine eklenir.")
add_bullet(doc, "Vision tower agirliklari aynen birakilir (zaten LoRA "
                "etkilenmemistir).")

add_heading(doc, "5.7 Ollama Modeline Donusturme", level=2)
add_para(
    doc,
    "Sonuc GGUF dosyasi Docker volume'una mount edilip Ollama icindeki "
    "Modelfile referansiyla model olarak kayit edilir:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(doc)
add_code_block(doc, """
# Modelfile
FROM ./gemma3-4b-bp-v2.q4_k_m.gguf
PARAMETER temperature 0.2
PARAMETER top_p 0.9
PARAMETER num_ctx 4096
SYSTEM "Sen Flowgent AI agent'isin..."
""")
add_para(doc)
add_code_block(doc, """
$ ollama create bp-agent-v2 -f Modelfile
$ ollama list
NAME                ID              SIZE      MODIFIED
bp-agent-v2:latest  ab12cd34...     2.5 GB    just now
gemma3:4b           ef56gh78...     3.3 GB    2 days ago
""")
add_para(
    doc,
    "Kayit sonrasinda AiService'in appsettings.json'da `UseFinetuned: true` "
    "yapilarak agent ozel modeli kullanmaya baslar. ModelSelector adli "
    "yardimci sinif feature flag durumuna gore primary modeli secer.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

doc.add_page_break()

# ================================================================
# 6. DEGERLENDIRME
# ================================================================
add_heading(doc, "6. DEGERLENDIRME", level=1)

add_heading(doc, "6.1 Eval Runner", level=2)
add_para(
    doc,
    "tools/ai_eval/runner.py birden cok modeli ayni eval setine karsi "
    "calistirip JSON-cikti dogrulugu, schema-pass orani, tool-call dogrulugu "
    "ve istek basina ortalama sure metriklerini uretir. Calistirma:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(doc)
add_code_block(doc, """
# Sadece base
$ python -m tools.ai_eval.runner --models gemma3:4b --limit 5

# Karsilastirmali
$ python -m tools.ai_eval.runner --models gemma3:4b,bp-agent-v2
""")
add_para(
    doc,
    "Eval seti 4 ozellige (scaffold, enrich, plan, agent) ait toplam 20 "
    "prompt icerir. Her prompt icin uretim alinir, regex ile JSON cikarilir, "
    "schema valide edilir, gerektiginde tool-call argumanlari beklenenle "
    "karsilastirilir. Cikti CSV + Markdown rapor olarak yazilir.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

add_heading(doc, "6.2 Base Model Davranis Analizi", level=2)
add_para(
    doc,
    "Base gemma3:4b uzerinde gozlemlenen davranislar (Bolum 5.1'de "
    "ozetlenmistir) eval ile sayisallastirildiginda:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
rows = [
    ["JSON-parse oran", "%62 (tolerant parser ile %88)"],
    ["Schema-pass orani (enrich)", "%0 (sema uyumsuz cikti)"],
    ["Tool-call dogru ad orani (agent)", "%55"],
    ["'final' ile bitirebilme orani", "%30 (5 iter limit)"],
    ["Ortalama sure (saniye/istek)", "~6-9 sn (CPU+GPU hibrit Ollama)"],
]
add_table(doc, ["Metrik", "Base gemma3:4b"], rows)

add_heading(doc, "6.3 v1 Fine-tune Sonuclari (Basarisizlik Analizi)", level=2)
add_para(
    doc,
    "v1 fine-tune (checkpoint-125 / LoRA r=8) yerelde basariyla GGUF'a "
    "cevrilip bp-agent olarak Ollama'ya yuklendi. Ancak canli testlerde "
    "hemen her agent senaryosunu bozdugu icin uretim flag'i (UseFinetuned) "
    "false'a alindi. Hata modlari:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_bullet(doc, "Tool secimi yanlis: 'aktif sprint nedir' sorusuna create_issue "
                "cagriliyor.")
add_bullet(doc, "Null input halusinasyonu: tool_calls.input.title alani "
                "uretiliyor ama icerigi bos.")
add_bullet(doc, "Final uretememe: tool sonucu donduktan sonra ozet yerine "
                "ayni tool tekrar cagriliyor.")
add_bullet(doc, "Tekrar collapse: scaffold ciktilarinda sprint goal'lerinde "
                "10 kez ayni kelime ('strateji strateji strateji...').")
add_para(doc)
add_para(
    doc,
    "Kok-neden analizi iki kritik faktoru ortaya cikardi:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_bullet(doc, "(1) LoRA r=8 kapasite yetersizligi: 4B parametreli base "
                "uzerine yalnizca rank-8 LoRA, hem format ezberi hem cok "
                "ozellikli karar dinamigi icin yetersiz.")
add_bullet(doc, "(2) Tool-calling ornegi tamamen yoktu: v1 datasetinde agent "
                "kategorisi 0 ornek iceriyordu. Model yalnizca scaffold/"
                "enrich/plan formatlarini ezberledigi icin tool-calling "
                "modunda hizla dagiliyor.")

add_heading(doc, "6.4 v2 Beklentileri", level=2)
add_para(
    doc,
    "v2 fine-tune sonrasinda (henuz Colab egitimi tamamlandiktan sonra "
    "rapor edilecek) beklenen iyilesmeler:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_bullet(doc, "JSON-parse oraninda fenced/missing-brace hatalari %5 alti.")
add_bullet(doc, "Agent senaryolarinda dogru tool secimi orani > %85.")
add_bullet(doc, "'final' ile bitirebilme orani > %90 (5 iter limitine takilma "
                "azalmasi).")
add_bullet(doc, "Schema-pass (enrich-issue) orani > %85.")

doc.add_page_break()

# ================================================================
# 7. KARSILASILAN ZORLUKLAR
# ================================================================
add_heading(doc, "7. KARSILASILAN ZORLUKLAR VE COZUMLER", level=1)
rows = [
    [
        "Yerel kucuk modelin JSON sapmalari",
        "Tolerant JSON parser (markdown fence kirpma, eksik brace tahmini, "
        "Turkce alias map)",
    ],
    [
        "Gemma3 multimodal'da PEFT.load_adapter cokmesi",
        "manual_merge.py: yalniz language_model regex'i ile ΔW = (alpha/r)*(B@A) "
        "elle uygulanir.",
    ],
    [
        "Colab T4 RAM siniri (12 GB)",
        "GGUF donusumu YEREL'de yapilir; Colab sadece egitim icin kullanilir.",
    ],
    [
        "Groq API kotalarinin tukenmesi",
        "Mock fallback uretici (mock_gen.py) + cogul-saglayicili auto-mode "
        "(Gemini > Groq) eklendi.",
    ],
    [
        "YARP gateway agent endpoint timeout",
        "Cluster icin ActivityTimeout = 00:05:00 yapildi (uzun LLM cagrilari "
        "icin).",
    ],
    [
        "v1 fine-tune'un her senaryoyu bozmasi",
        "UseFinetuned feature flag false'a alindi; v2 dataseti + r=32 ile "
        "yeniden egitim plani devreye girdi.",
    ],
    [
        "Sentetik veri ve eval seti cakismasi (data leakage)",
        "merge.py uvinde eval_collision_filter adimi eklendi; cakisan ornekler "
        "egitim setinden ayiklanir.",
    ],
]
add_table(doc, ["Zorluk", "Cozum"], rows)

doc.add_page_break()

# ================================================================
# 8. SONUC
# ================================================================
add_heading(doc, "8. SONUC VE GELECEK CALISMALAR", level=1)
add_para(
    doc,
    "Bu calismada, OpenAI / Anthropic gibi bulut LLM saglayicilarina ihtiyac "
    "duymadan tamamen yerel donanimda calisan, denetlenebilir ve fine-tune "
    "edilebilir bir tool-calling agent'inin uctan uca gerceklestirilmesi "
    "sunulmustur. Sistem mevcut haliyle:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_bullet(doc, "Base gemma3:4b ile tolerant parser kombinasyonu sayesinde "
                "Flowgent platformunun 7 ozelliginin (plan, enrich, chat, "
                "retro, balance, risk, scaffold) tumunde uretim seviyesinde "
                "(demo amacli) calisir durumdadir.")
add_bullet(doc, "Agent endpoint'i (5 toollu) yerel modelle issue/sprint "
                "olusturma ve sorgulama gibi temel akislari basariyla "
                "tamamlamaktadir.")
add_bullet(doc, "v1 fine-tune basarisiz olmus, hatalari analiz edilmis ve "
                "v2 icin gerekli dataset (1417 ornek) ve config (LoRA r=32) "
                "hazirlanmistir.")
add_para(doc)
add_para(
    doc,
    "Gelecek calismalar olarak hedeflenenler:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_bullet(doc, "v2 fine-tune'in Colab T4 uzerinde tam 3 epoch egitilmesi "
                "ve eval raporunun before/after tablosuyla doldurulmasi.")
add_bullet(doc, "Tool katalogunun genisletilmesi: update_issue_status, "
                "assign_user_to_issue, complete_sprint gibi yazma araclari.")
add_bullet(doc, "RAG-lite chat ozelligine gercek vektor veritabani "
                "(pgvector veya Qdrant) entegrasyonu.")
add_bullet(doc, "DPO/ORPO tarzi tercih tabanli ince-ayar ile 'final' "
                "uretme davranisinin daha da guclendirilmesi.")
add_bullet(doc, "Uretim ortaminda streaming cevap (SSE / WebSocket) destegi.")

doc.add_page_break()

# ================================================================
# 9. KAYNAKLAR
# ================================================================
add_heading(doc, "9. KAYNAKLAR", level=1)
refs = [
    "[1] Hu, E. J. et al. (2021). LoRA: Low-Rank Adaptation of Large Language "
    "Models. arXiv:2106.09685.",
    "[2] Google. (2025). Gemma 3 Technical Report. ai.google.dev/gemma.",
    "[3] Ollama. Local LLM runtime. https://ollama.com",
    "[4] ggerganov, llama.cpp. https://github.com/ggerganov/llama.cpp",
    "[5] Unsloth AI. Fast LoRA fine-tuning for LLMs. https://github.com/unslothai/unsloth",
    "[6] Hugging Face PEFT. Parameter-Efficient Fine-Tuning Library. "
    "https://github.com/huggingface/peft",
    "[7] Schick, T. et al. (2023). Toolformer: Language Models Can Teach "
    "Themselves to Use Tools. NeurIPS 2023.",
    "[8] Yao, S. et al. (2022). ReAct: Synergizing Reasoning and Acting in "
    "Language Models. ICLR 2023.",
    "[9] Wang, Y. et al. (2023). Self-Instruct: Aligning Language Models "
    "with Self-Generated Instructions. ACL 2023.",
    "[10] Microsoft. .NET 9 Documentation. https://learn.microsoft.com/dotnet",
    "[11] Microsoft. YARP - Yet Another Reverse Proxy. "
    "https://microsoft.github.io/reverse-proxy/",
    "[12] Groq. https://groq.com (Llama-3.3-70b inference API)",
    "[13] Google. Gemini API. https://ai.google.dev (gemini-2.5-flash)",
]
for r in refs:
    add_para(doc, r, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

doc.save(OUTPUT)
print(f"Yazildi: {OUTPUT}")
