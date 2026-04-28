# -*- coding: utf-8 -*-
"""Generate the interim report (Ara Rapor) for the BitirmeProject in the same
layout as docs/report/AppRay_Ara_Raporu_222201310.docx."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

OUTPUT = r"docs/report/BitirmeProject_Ara_Raporu.docx"

PROJECT_TITLE_TR = (
    "YAPAY ZEKA DESTEKLİ MİKROSERVİS MİMARİLİ ÇEVİK PROJE YÖNETİM PLATFORMU"
)
PROJECT_TITLE_EN = (
    "AI-ASSISTED MICROSERVICE-BASED AGILE PROJECT MANAGEMENT PLATFORM"
)
STUDENT_NO = "XXXXXXXXX"
STUDENT_NAME = "Metin KARYAĞDI"
ADVISOR = "Dr. Öğr. Üyesi [Danışman Adı]"
MONTH_TR = "Nisan 2026"
MONTH_EN = "April 2026"
CITY = "AMASYA"
UNIVERSITY_TR = "T.C. AMASYA ÜNİVERSİTESİ"
UNIVERSITY_EN = "T.C. AMASYA UNIVERSITY"
FACULTY_TR = "MÜHENDİSLİK FAKÜLTESİ"
FACULTY_EN = "FACULTY OF ENGINEERING"
DEPT_TR = "BİLGİSAYAR MÜHENDİSLİĞİ BÖLÜMÜ"


def add_para(doc, text, *, style=None, bold=False, align=None, size=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        if bold:
            run.bold = True
        if size:
            run.font.size = Pt(size)
    return p


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    return h


def add_table(doc, header, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Table Grid"
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            table.rows[ri].cells[ci].text = val
    return table


doc = Document()

# ============= KAPAK =============
for _ in range(3):
    add_para(doc, "")
add_para(doc, UNIVERSITY_TR, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para(doc, FACULTY_TR, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para(doc, DEPT_TR, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para(doc, "")
add_para(doc, "")
add_para(doc, "BİTİRME PROJESİ ARA RAPORU", bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
add_para(doc, "")
add_para(doc, "")
add_para(doc, PROJECT_TITLE_TR, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
for _ in range(4):
    add_para(doc, "")
add_para(doc, f"{STUDENT_NO} - {STUDENT_NAME}",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para(doc, f"Danışman: {ADVISOR}",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
for _ in range(5):
    add_para(doc, "")
add_para(doc, MONTH_TR, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para(doc, CITY, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

doc.add_page_break()

# ============= ÖZET =============
add_para(doc, UNIVERSITY_TR, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
add_para(doc, FACULTY_TR, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
add_para(doc, MONTH_TR, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, PROJECT_TITLE_TR, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
add_para(doc, f"({STUDENT_NO} {STUDENT_NAME})",
         align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_heading("ÖZET", level=0)
add_para(doc,
    "Bu çalışmada yazılım ekiplerinin proje, sprint ve görev yönetimi süreçlerini "
    "tek bir panelde toplayan, yapay zeka destekli ve mikroservis mimarisi üzerine "
    "kurulmuş bir platformun geliştirilmesi amaçlanmıştır. Proje kapsamında "
    "kullanıcı kimliği, organizasyon yönetimi, davet sistemi, proje açma, sprint "
    "planlama, issue (görev) takibi, dosya eki yükleme, gerçek zamanlı bildirim "
    "ve yapay zeka tabanlı yardımcı özellikler birbirinden bağımsız servisler "
    "şeklinde ele alınmıştır. Sistem; IdentityService, ProjectService, "
    "SprintService, IssueService, NotificationService, StorageService, AiService "
    "ve BFF/Gateway bileşenleri üzerine kurulmuş olup PostgreSQL, RabbitMQ, "
    "Redis, MailHog, Seq ve Ollama altyapılarından yararlanmaktadır. Ara rapor "
    "aşamasında kullanıcı kayıt/giriş, JWT tabanlı kimlik doğrulama, "
    "organizasyon davet akışı, Kanban görünümlü görev panosu, sprint takibi, "
    "SignalR üzerinden gerçek zamanlı bildirimler, dosya eki yükleme ve yedi "
    "farklı yapay zeka özelliği (plan üretimi, sohbet, görev zenginleştirme, "
    "sprint risk analizi, yük dengeleme önerisi, retrospektif ve araç çağırabilen "
    "asistan) çalışır hale getirilmiştir. Ayrıca yerel olarak çalışan "
    "gemma3:4b modelinin LoRA tabanlı ince ayarı için yedi fazlı bir veri "
    "toplama ve eğitim hattı kurulmuş, sentetik veri üretim araçları ile "
    "değerlendirme altyapısı geliştirilmiştir. Mevcut durum, projenin bitirme "
    "aşamasına kadar daha kapsamlı testler, ince ayarlı model entegrasyonu ve "
    "kullanıcı deneyimi iyileştirmeleri ile olgunlaştırılabileceğini "
    "göstermektedir."
)
add_para(doc,
    "Anahtar Kelimeler: Proje yönetimi, mikroservis, çevik yazılım, yapay zeka, "
    "büyük dil modeli, ince ayar, gemma3, RabbitMQ, SignalR")

# ============= ABSTRACT =============
add_para(doc, "")
add_para(doc, UNIVERSITY_EN, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
add_para(doc, FACULTY_EN, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
add_para(doc, MONTH_EN, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, PROJECT_TITLE_EN, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
add_para(doc, f"({STUDENT_NO} {STUDENT_NAME})",
         align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_heading("ABSTRACT", level=0)
add_para(doc,
    "This study aims to develop an AI-assisted, microservice-based platform that "
    "consolidates the project, sprint and task management workflows of software "
    "teams into a single panel. The platform handles user identity, organization "
    "management, invitations, project creation, sprint planning, issue tracking, "
    "attachment uploads, real-time notifications and AI-assisted features as "
    "independent services. The system is built on IdentityService, "
    "ProjectService, SprintService, IssueService, NotificationService, "
    "StorageService, AiService and BFF/Gateway components, while leveraging "
    "PostgreSQL, RabbitMQ, Redis, MailHog, Seq and Ollama infrastructures. At "
    "the interim report stage, user registration/login, JWT-based "
    "authentication, the organization invite flow, a Kanban-style issue board, "
    "sprint tracking, real-time notifications via SignalR, attachment uploads "
    "and seven AI features (plan generation, chat, issue enrichment, sprint "
    "risk analysis, workload balancing, retrospective and a tool-calling "
    "assistant) have been implemented. In addition, a seven-phase data "
    "collection and training pipeline has been established for the LoRA-based "
    "fine-tuning of the locally running gemma3:4b model, including synthetic "
    "data generation tools and an evaluation framework. The current state "
    "shows that the project has reached a functional intermediate level and "
    "can be matured until the final report through broader testing, the "
    "integration of the fine-tuned model, and user experience improvements."
)
add_para(doc,
    "Keywords: Project management, microservices, agile, artificial "
    "intelligence, large language model, fine-tuning, gemma3, RabbitMQ, "
    "SignalR")

# ============= TEŞEKKÜR =============
doc.add_heading("TEŞEKKÜR", level=0)
add_para(doc,
    "Bu çalışmanın ara rapor aşamasına gelmesinde yönlendirmeleri ve "
    "değerlendirmeleriyle katkı sağlayan danışmanım " + ADVISOR + "'a teşekkür "
    "ederim. Ayrıca kullanılan açık kaynak kütüphaneler, .NET ve React "
    "ekosistemleri, Ollama topluluğu ile dokümantasyonların hazırlanmasında "
    "katkısı bulunan tüm topluluklara da teşekkür ederim."
)

# ============= İÇİNDEKİLER =============
doc.add_heading("İÇİNDEKİLER", level=0)
toc_items = [
    ("        Sayfa", ""),
    ("ÖZET", "i"),
    ("ABSTRACT", "ii"),
    ("TEŞEKKÜR", "iii"),
    ("İÇİNDEKİLER", "iv"),
    ("ÇİZELGELERİN LİSTESİ", "v"),
    ("ŞEKİLLERİN LİSTESİ", "vi"),
    ("SİMGELER VE KISALTMALAR", "vii"),
    ("1. GİRİŞ", "1"),
    ("2. LİTERATÜR TARAMASI", "1"),
    ("2.1. Çevik proje yönetimi ve sprint tabanlı planlama", "1"),
    ("2.2. Mikroservis mimarisi ve event-driven yaklaşım", "2"),
    ("2.3. Büyük dil modelleri ve yerel çalışan LLM'ler", "2"),
    ("2.4. Tool-calling ve asistan tabanlı kullanım senaryoları", "2"),
    ("3. YÖNTEM VE SİSTEM MİMARİSİ", "2"),
    ("3.1. Genel mimari yaklaşım", "2"),
    ("3.2. Kimlik, organizasyon ve davet akışı", "3"),
    ("3.3. Proje, sprint ve issue veri akışı", "3"),
    ("3.4. Yapay zeka katmanı ve ince ayar hattı", "3"),
    ("3.5. Kullanılan servisler ve görevleri", "3"),
    ("4. GELİŞTİRME SÜRECİ, MEVCUT DURUM VE GÖRSEL ÇIKTILAR", "4"),
    ("4.1. Geliştirme sürecinin adım adım ilerleyişi", "4"),
    ("4.2. Güncel arayüz ve çalışma örnekleri", "7"),
    ("5. FİNALE KADAR PLANLANAN GELİŞTİRMELER VE EKSİKLER", "12"),
    ("5.1. Bu aşamada tamamlanan temel işlevler", "12"),
    ("5.2. Henüz tamamlanmamış veya güçlendirilmesi gereken noktalar", "12"),
    ("5.3. Finale kadar planlanan somut adımlar", "13"),
    ("6. SONUÇ VE ÖNERİLER", "13"),
    ("7. KAYNAKLAR", "14"),
    ("8. EKLER", "15"),
    ("8.1. Proje yapısı", "15"),
    ("9. ÖZGEÇMİŞ", "16"),
]
for label, page in toc_items:
    if page:
        add_para(doc, f"{label}\t\t{page}")
    else:
        add_para(doc, label)

# ============= ÇİZELGELER =============
doc.add_heading("ÇİZELGELERİN LİSTESİ", level=0)
add_para(doc, "Çizelge 3.1. Platform servisleri ve görevleri\t\t12")
add_para(doc, "Çizelge 4.1. Geliştirme adımları ve mevcut durum özeti\t\t14")

# ============= ŞEKİLLER =============
doc.add_heading("ŞEKİLLERİN LİSTESİ", level=0)
sekil_list = [
    "Şekil 4.1. Giriş ve kayıt ekranı",
    "Şekil 4.2. Organizasyon onboarding ve davet akışı",
    "Şekil 4.3. Proje listesi ve oluşturma ekranı",
    "Şekil 4.4. Sprint planlama ve takip ekranı",
    "Şekil 4.5. Kanban görünümlü issue panosu",
    "Şekil 4.6. Issue detay paneli ve AI zenginleştirme",
    "Şekil 4.7. AI Asistan — yeni proje taslağı (scaffold) ekranı",
    "Şekil 4.8. AI Asistan — araç çağıran asistan (agent) ekranı",
    "Şekil 4.9. SignalR tabanlı bildirim merkezi",
    "Şekil 4.10. Docker servislerinin çalışır durumu",
    "Şekil 4.11. Seq üzerinden toplanan log akışı",
    "Şekil 4.12. RabbitMQ kuyruk yönetim ekranı",
    "Şekil 4.13. Sentetik veri üretim ilerleme görseli",
]
for s in sekil_list:
    add_para(doc, f"{s}\t\t{15 + sekil_list.index(s)}")

# ============= KISALTMALAR =============
doc.add_heading("SİMGELER VE KISALTMALAR", level=0)
kisaltmalar = [
    ("Kısaltmalar", "Açıklamalar"),
    ("AI", "Artificial Intelligence (Yapay Zeka)"),
    ("API", "Application Programming Interface"),
    ("BFF", "Backend For Frontend"),
    ("CRUD", "Create, Read, Update, Delete"),
    ("DTO", "Data Transfer Object"),
    ("EF", "Entity Framework"),
    ("JWT", "JSON Web Token"),
    ("LLM", "Large Language Model (Büyük Dil Modeli)"),
    ("LoRA", "Low-Rank Adaptation"),
    ("ORM", "Object-Relational Mapping"),
    ("REST", "Representational State Transfer"),
    ("SaaS", "Software as a Service"),
    ("SDK", "Software Development Kit"),
    ("SMTP", "Simple Mail Transfer Protocol"),
    ("UI", "User Interface"),
]
for k, a in kisaltmalar:
    add_para(doc, f"{k}\t{a}")

doc.add_page_break()

# ============= 1. GİRİŞ =============
doc.add_heading("1. GİRİŞ", level=1)
add_para(doc,
    "Yazılım projelerinin gün geçtikçe artan karmaşıklığı, ekiplerin "
    "geliştirme sürecini disiplinli biçimde yönetmesini zorunlu kılmaktadır. "
    "Çevik (agile) yaklaşımların yaygınlaşmasıyla birlikte sprint tabanlı "
    "planlama, görev (issue) takibi, retrospektif değerlendirmeler ve takım "
    "içi şeffaflık temel ihtiyaçlar haline gelmiştir. Pratikte ise bu "
    "süreçlerde kullanılan araçlar genellikle birbirinden bağımsız çalışır; "
    "kullanıcı, projeyi bir araçta planlar, görevleri ikinci bir araçta "
    "takip eder, raporlamayı üçüncü bir araçtan alır. Bu durum hem "
    "operasyonel zaman kaybına yol açmakta hem de tek bir projenin "
    "bütünleşik resminin kaybolmasına sebep olmaktadır."
)
add_para(doc,
    "Bu proje söz konusu dağınıklığı azaltmak amacıyla, yazılım ekiplerinin "
    "organizasyon, proje, sprint ve görev süreçlerini tek bir panelde "
    "toplayan, mikroservis mimarisi üzerine kurulmuş ve yapay zeka destekli "
    "bir platform olarak tasarlanmıştır. Kullanıcı, bir organizasyon "
    "oluşturur veya davet aldığı bir organizasyona katılır; ardından "
    "projeler açar, sprint planlar, görevleri Kanban panosu üzerinden "
    "takip eder. Yapay zeka katmanı ise yalnızca metinsel bir sohbet "
    "değil; proje taslağı çıkarma, sprint risk analizi, yük dengeleme "
    "önerisi, retrospektif yorumlama ve doğrudan veri tabanına yazabilen "
    "araç çağıran (tool-calling) bir asistan biçiminde devreye girmektedir."
)
add_para(doc,
    "Proje seçim sürecinde üç temel etken belirleyici olmuştur. Birincisi, "
    ".NET tabanlı mikroservislerle gerçek bir SaaS uygulaması geliştirme "
    "isteğidir. İkincisi, büyük dil modellerinin (LLM) yerel olarak "
    "çalıştırılması ve LoRA tabanlı ince ayar ile özelleştirilmesinin "
    "öğretici, özgün ve güncel bir mühendislik problemi sunmasıdır. "
    "Üçüncüsü ise projenin yapılabilirliğidir; çünkü Docker tabanlı "
    "servisler, açık kaynak araçlar ve Ollama gibi yerel LLM altyapıları "
    "sayesinde proje adım adım geliştirilebilir bir çerçeveye "
    "oturtulabilmiştir."
)
add_para(doc,
    "Ara raporun amacı, projenin başlangıcından itibaren izlenen geliştirme "
    "yolunu, şu ana kadar yapılan çalışmaları, mevcut durumda çalışan "
    "modülleri ve finale kadar tamamlanması planlanan eksikleri sistematik "
    "biçimde ortaya koymaktır. Bu doğrultuda raporda önce konuya ilişkin "
    "kısa bir literatür özeti verilmiş, daha sonra sistem mimarisi "
    "açıklanmış, geliştirme süreci adım adım anlatılmış ve elde edilen "
    "güncel arayüz ile altyapı çıktıları görsellerle desteklenmiştir."
)

# ============= 2. LİTERATÜR =============
doc.add_heading("2. LİTERATÜR TARAMASI", level=1)

doc.add_heading("2.1. Çevik proje yönetimi ve sprint tabanlı planlama", level=2)
add_para(doc,
    "Çevik yazılım geliştirme yaklaşımı, kısa iterasyonlar (sprint) "
    "üzerinden sürekli teslim, sürekli geri bildirim ve değişime açık "
    "planlama ilkelerine dayanmaktadır. Scrum ve Kanban gibi çerçeveler, "
    "iş öğelerinin (issue) belirli durumlar arasında akmasını ve sprint "
    "sınırları içinde tamamlanmasını öngörür [1]. Akademik çalışmalar ve "
    "endüstri raporları, çevik araçların ekip içi şeffaflığı artırdığını ve "
    "tahminleme kalitesini iyileştirdiğini göstermektedir [2]. Ancak "
    "popüler ticari araçlar (Jira, Trello, Asana) tek başına çalıştığında, "
    "ekipler genellikle dokümantasyonu, raporlamayı ve yapay zeka destekli "
    "yardımı farklı eklenti veya ürünlerle sağlamak zorunda kalır."
)

doc.add_heading("2.2. Mikroservis mimarisi ve event-driven yaklaşım", level=2)
add_para(doc,
    "Mikroservis mimarisi, monolitik uygulamaların ölçeklenme ve bakım "
    "zorluklarına karşı geliştirilmiş bir yaklaşımdır. Her servis kendi "
    "veri tabanına ve sorumluluk alanına sahiptir ve servisler arasındaki "
    "iletişim çoğunlukla REST ya da mesajlaşma altyapıları üzerinden "
    "sağlanır [3]. RabbitMQ gibi mesaj kuyrukları, servisler arasında "
    "gevşek bağlılık (loose coupling) ve dayanıklılık sağlar [4]. Bu "
    "projede de Outbox deseni ile birlikte RabbitMQ kullanılarak "
    "servisler arası bildirim ve olay yayını gerçekleştirilmiştir."
)

doc.add_heading("2.3. Büyük dil modelleri ve yerel çalışan LLM'ler", level=2)
add_para(doc,
    "Son yıllarda büyük dil modelleri (LLM) yazılım geliştirme süreçlerinde "
    "kod üretimi, dokümantasyon ve doğal dil arayüzleri için yoğun biçimde "
    "kullanılmaktadır [5]. Bulut tabanlı modellerin (GPT, Claude, Gemini) "
    "yanı sıra, daha küçük parametre sayısına sahip ve yerel olarak "
    "çalıştırılabilen modeller (Llama, Gemma, Qwen) gizlilik, maliyet ve "
    "çevrim dışı kullanım gerekçeleriyle giderek tercih edilmektedir [6]. "
    "Ollama gibi araçlar bu yerel modellerin Docker ortamında "
    "çalıştırılmasını kolaylaştırmaktadır [7]. Bu projede gemma3:4b "
    "modeli temel alınarak, LoRA tabanlı ince ayar (PEFT) ile alana özgü "
    "davranışların kazandırılması hedeflenmiştir [8]."
)

doc.add_heading("2.4. Tool-calling ve asistan tabanlı kullanım senaryoları",
                level=2)
add_para(doc,
    "Geleneksel sohbet arayüzlerinin ötesinde, modellerin önceden "
    "tanımlanmış araçları (function calling / tool use) çağırarak gerçek "
    "iş süreçlerine müdahale edebilmesi, üretkenlik açısından kritik bir "
    "yetenek olarak öne çıkmıştır [9]. Bu yaklaşımda model, kullanıcının "
    "doğal dil ifadesini analiz eder, hangi aracı hangi parametrelerle "
    "çağıracağını belirler ve sonucu kullanıcıya doğal dil ile aktarır. "
    "Projede de bu yaklaşım benimsenmiş; CreateIssue, CreateSprint, "
    "AddIssueToSprint, GetActiveSprint ve GetProjectIssues olmak üzere "
    "beş araç tanımlanmış ve modelin bu araçları döngüsel biçimde "
    "kullanabildiği bir AgentLoop yapısı kurulmuştur."
)

# ============= 3. MİMARİ =============
doc.add_heading("3. YÖNTEM VE SİSTEM MİMARİSİ", level=1)

doc.add_heading("3.1. Genel mimari yaklaşım", level=2)
add_para(doc,
    "Sistem, her birinin ayrı veri tabanına ve sorumluluk alanına sahip "
    "olduğu mikroservislerden oluşmaktadır. Frontend tarafında React + "
    "TypeScript + Vite tabanlı tek sayfa uygulama (SPA), arka uçta ise "
    "ASP.NET Core 8/9 tabanlı sekiz farklı servis bulunur: IdentityService, "
    "ProjectService, SprintService, IssueService, NotificationService, "
    "StorageService, AiService ve BFF/Gateway katmanları. Tüm servisler "
    "Docker Compose üzerinden ayağa kaldırılır; ortak altyapı bileşenleri "
    "olarak PostgreSQL 16, RabbitMQ 3.13, Redis 7, MailHog, Seq ve Ollama "
    "(gemma3:4b) kullanılmıştır."
)
add_para(doc,
    "Servisler arası senkron iletişim REST üzerinden, asenkron iletişim "
    "ise RabbitMQ üzerinden gerçekleştirilir. Bildirim servisi SignalR ile "
    "frontend'e gerçek zamanlı sinyal gönderir. AI servisi Ollama HTTP "
    "API'si ile yerel modeli çağırır; üretim ortamına geçildiğinde "
    "feature flag (Ollama:UseFinetuned) ile ince ayarlı modele geçiş tek "
    "satır konfigürasyon değişikliği ile yapılabilmektedir."
)

doc.add_heading("3.2. Kimlik, organizasyon ve davet akışı", level=2)
add_para(doc,
    "Kullanıcı doğrulaması IdentityService tarafından JWT bazlı olarak "
    "yapılır. Token üretimi sırasında kullanıcının aktif organizasyonu ve "
    "rolü (Admin, Manager, Member) claim olarak token içine yerleştirilir. "
    "Bu sayede diğer servisler (Project, Sprint, Issue) gelen istekte "
    "yetki kontrolünü token'dan hızlıca yapabilir; ayrı bir yetki sorgusu "
    "yapmak gerekmez. Yeni bir kullanıcının sisteme dahil olması, "
    "kayıt → onay → organizasyon seçme veya davet kabul etme adımlarından "
    "oluşmaktadır. Davet maillerinin lokal geliştirme ortamında test "
    "edilebilmesi için MailHog SMTP yakalayıcısı entegre edilmiştir."
)

doc.add_heading("3.3. Proje, sprint ve issue veri akışı", level=2)
add_para(doc,
    "Bir kullanıcı, ait olduğu organizasyonda yeni bir proje oluşturduğunda "
    "ProjectService bu kaydı kendi veri tabanına yazar ve bir "
    "OrganizationId denormalize bilgisini issue ve sprint kayıtlarına da "
    "yansıtır. Bu denormalizasyon, sonradan eklenen yetki sıkılaştırmaları "
    "kapsamında controller seviyesinde fast-fail kontrollerinin yapılmasını "
    "mümkün kılmıştır. Issue'ler Open / InProgress / Done durumları arasında "
    "Kanban panosu üzerinden sürüklenip bırakılarak hareket ettirilir; "
    "her hareket NotificationService'e bir olay olarak yayınlanır."
)
add_para(doc,
    "Sprint'ler tanımlı bir zaman aralığına sahiptir. Aktif sprint "
    "üzerindeki ilerleme, AI servisinin sprint-risk ve suggest-balance "
    "uçlarında girdi olarak kullanılır. Tamamlanmış sprintler için "
    "retrospektif uç noktası, modelden geri dönüş, iyileştirme alanları "
    "ve takım gözlemi şeklinde yapılandırılmış bir çıktı üretir."
)

doc.add_heading("3.4. Yapay zeka katmanı ve ince ayar hattı", level=2)
add_para(doc,
    "Yapay zeka katmanı iki ana bileşene sahiptir. Birincisi, çalışma "
    "zamanında çağrılan AiService; ikincisi ise model davranışını "
    "iyileştirmek üzere kurulan ince ayar (fine-tune) hattıdır. "
    "AiService üzerinde toplam yedi farklı uç nokta vardır: chat, "
    "generate-plan, enrich-issue, sprint-risk, suggest-balance, "
    "retrospective ve agent. Agent uç noktası, AgentLoop adı verilen "
    "yapı ile çalışır; modelin döndürdüğü JSON üzerinden tool_calls "
    "alanı ayrıştırılır, ilgili araç çalıştırılır, sonucu modele geri "
    "verilir ve final cevaba ulaşana kadar bu döngü maksimum beş tur "
    "ile sınırlandırılır. Her tool çağrısı AiToolExecution tablosuna "
    "denetim (audit) amacıyla kaydedilir."
)
add_para(doc,
    "İnce ayar hattı yedi fazdan oluşmaktadır. Faz 0 (hazırlık), Faz 1 "
    "(veri toplama), Faz 2 (eğitim pipeline), Faz 3 (eval runner), Faz 4 "
    "(deploy altyapısı), Faz 5 (tool-calling) ve Faz 6 (savunma raporu "
    "şablonu). Faz 1 kapsamında Groq API üzerinden llama-3.3-70b "
    "modeline çağrı yapılarak, üç AI özelliği (scaffold-project, "
    "enrich-issue, generate-plan) için sentetik girdi-çıktı çiftleri "
    "üretilmektedir. Hedef en az dokuz yüz örnektir; ara rapor tarihi "
    "itibarıyla yaklaşık dört yüz örnek toplanmış olup üretim devam "
    "etmektedir. Eğitim, Google Colab üzerinde LoRA tabanlı PEFT "
    "kütüphanesi ile gerçekleştirilecek, ortaya çıkan adapter Ollama "
    "Modelfile aracılığıyla bp-agent isimli bir model olarak "
    "kullanılacaktır."
)

doc.add_heading("3.5. Kullanılan servisler ve görevleri", level=2)
add_para(doc, "Çizelge 3.1. Platform servisleri ve görevleri")
add_table(doc,
    ["Servis", "Görev", "Mevcut Durum"],
    [
        ["IdentityService.Api",
         "Kullanıcı kayıt/giriş, JWT, organizasyon, davet ve rol yönetimi",
         "Çalışır durumda"],
        ["ProjectService.Api",
         "Proje CRUD, organizasyon-proje ilişkisi, üyelik kontrolü",
         "Çalışır durumda"],
        ["SprintService.Api",
         "Sprint CRUD, aktif sprint takibi, sprint-issue ilişkisi",
         "Çalışır durumda"],
        ["IssueService.Api",
         "Issue CRUD, durum geçişleri (Open/InProgress/Done), atama",
         "Çalışır durumda"],
        ["NotificationService.Api",
         "Olay tüketimi, SignalR yayını, kullanıcı bildirim merkezi",
         "Çalışır durumda"],
        ["StorageService.Api",
         "Dosya/ek yükleme, indirme ve issue ile ilişkilendirme",
         "Çalışır durumda"],
        ["AiService.Api",
         "Yedi yapay zeka uç noktası, AgentLoop, tool dispatch ve audit",
         "Çalışır durumda"],
        ["Bff.Api / Gateway",
         "Frontend için aggregator uç noktaları ve kapı görevi",
         "Çalışır durumda"],
        ["PostgreSQL 16",
         "Her servis için ayrı veri tabanı (identity-db, project-db, ...)",
         "Çalışır durumda"],
        ["RabbitMQ 3.13",
         "Servisler arası asenkron olay yayını, Outbox tüketimi",
         "Çalışır durumda"],
        ["Redis 7",
         "Önbellek altyapısı (henüz aktif kullanım yok, ileri faz)",
         "Hazır, kullanılmıyor"],
        ["MailHog",
         "Lokal SMTP yakalayıcı, davet ve doğrulama maillerinin testi",
         "Çalışır durumda"],
        ["Seq",
         "Yapılandırılmış log toplayıcı, servisler arası izleme",
         "Çalışır durumda"],
        ["Ollama (gemma3:4b)",
         "Yerel LLM sunucusu, AI servisinin model arka ucu",
         "Çalışır durumda"],
    ]
)
add_para(doc,
    "Bu yapı sayesinde kullanıcı tarafında görünen tek arayüz korunurken, "
    "arka planda farklı görevler birbirinden ayrılmış servislerde "
    "yürütülmektedir. Böylece hem bakım kolaylığı hem de geliştirme "
    "aşamalarının kontrollü şekilde ilerletilmesi mümkün olmuştur."
)

# ============= 4. GELİŞTİRME =============
doc.add_heading("4. GELİŞTİRME SÜRECİ, MEVCUT DURUM VE GÖRSEL ÇIKTILAR",
                level=1)

doc.add_heading("4.1. Geliştirme sürecinin adım adım ilerleyişi", level=2)
add_para(doc,
    "Projeye başlanırken ilk olarak çalışmanın kapsamı belirlenmiştir. "
    "Amaç, çevik yazılım ekiplerinin organizasyon, proje, sprint ve görev "
    "süreçlerini tek panelde yöneten ve yapay zeka destekli bir platform "
    "kurmaktır. Bu aşamada hangi servislerin ayrı tutulacağı, servisler "
    "arası iletişimin senkron mu yoksa asenkron mu olacağı, hangi "
    "yapay zeka özelliklerinin geliştirileceği ve kullanıcı akışının "
    "nasıl olacağı tasarlanmıştır."
)
add_para(doc,
    "İkinci aşamada Docker Compose tabanlı altyapı kurulmuştur. "
    "PostgreSQL 16, RabbitMQ 3.13, Redis 7, MailHog, Seq ve Ollama "
    "servisleri konteyner olarak ayağa kaldırılmış; her mikroservis için "
    "ayrı veri tabanı tanımlanmıştır. Geliştirici ortamında kullanılan "
    ".env.example dosyası ile gizli anahtarların ortam değişkenlerinden "
    "okunması sağlanmıştır."
)
add_para(doc,
    "Üçüncü aşamada IdentityService geliştirilmiştir. Kullanıcı kayıt, "
    "giriş, parola hashleme (BCrypt), JWT üretimi ve doğrulaması, e-posta "
    "doğrulama akışı ve admin için varsayılan kullanıcıyı tohumlayan "
    "AdminUserSeeder bu aşamada eklenmiştir. Sonraki adımda organizasyon "
    "kavramı sisteme dahil edilmiş; kullanıcının aktif organizasyonu "
    "(LastActiveOrg) JWT içinde claim olarak yer almaya başlamıştır."
)
add_para(doc,
    "Dördüncü aşamada ProjectService, SprintService ve IssueService "
    "geliştirilmiştir. Her servis kendi veri tabanına sahip biçimde "
    "tasarlanmış; servisler arası tutarlılık denormalizasyon ve event "
    "yayını ile sağlanmıştır. Issue durum geçişleri için event yayını "
    "RabbitMQ üzerinden gerçekleştirilmiş, NotificationService bu "
    "olayları dinleyerek ilgili kullanıcılara SignalR üzerinden gerçek "
    "zamanlı bildirim göndermiştir."
)
add_para(doc,
    "Beşinci aşamada organizasyon davet sistemi tamamlanmıştır. "
    "Yetkili kullanıcı, e-posta üzerinden davet gönderebilmekte; davetli "
    "kullanıcı maildeki bağlantı ile sisteme katılabilmektedir. Bu akışın "
    "uçtan uca test edilebilmesi için MailHog yakalayıcısı kullanılmıştır."
)
add_para(doc,
    "Altıncı aşamada AiService geliştirilmiştir. Bu servis Ollama HTTP "
    "API'si ile gemma3:4b modelini çağırmakta ve toplam yedi farklı uç "
    "nokta sunmaktadır: chat, generate-plan, enrich-issue, sprint-risk, "
    "suggest-balance, retrospective ve agent. Agent uç noktasında "
    "AgentLoop yapısı kurulmuş, modelin JSON çıktısının eksik kapanan "
    "süslü parantezleri ve markdown fence sarması gibi davranışlarına "
    "karşı toleranslı bir parser eklenmiştir."
)
add_para(doc,
    "Yedinci aşamada AI Asistan sayfası geliştirilmiştir. Bu sayfa iki "
    "moda sahiptir: Mevcut bir projede araç çağıran asistan modu (agent) "
    "ve yeni proje taslağı çıkaran scaffold modu. Scaffold modunda "
    "kullanıcı, doğal dilde proje tanımı yazar; sistem önce taslağı "
    "(proje + sprintler + issuelar) üretir, kullanıcı onayladığında "
    "ilgili kayıtları gerçek servislere yazar. Concurrency hatasına "
    "yol açan iki ayrı SaveChangesAsync çağrısı tek bir Add+Save "
    "işlemine indirilmiştir."
)
add_para(doc,
    "Sekizinci aşamada modelin LoRA tabanlı ince ayarı için altyapı "
    "kurulmuştur. tools/ai_data_collector altında sentetik veri üretim "
    "araçları, Groq ve Gemini sağlayıcıları ile birlikte hazırlanmıştır. "
    "tools/ai-finetune altında ise eğitim hazırlık scripti, "
    "tools/ai_eval altında değerlendirme runner'ı bulunmaktadır. Faz 1 "
    "kapsamında üç AI özelliği için yaklaşık dokuz yüz hedefli sentetik "
    "örneğin üretimi devam etmektedir."
)
add_para(doc,
    "Dokuzuncu aşamada admin paneli, dosya eki yükleme, bildirim merkezi "
    "ve onboarding gibi tamamlayıcı modüller geliştirilmiştir. Bu "
    "aşamaya kadar yapılan çalışmaların özeti Çizelge 4.1'de "
    "verilmiştir."
)

add_para(doc, "Çizelge 4.1. Geliştirme adımları ve mevcut durum özeti")
add_table(doc,
    ["Aşama", "Yapılanlar", "Durum"],
    [
        ["Problem ve kapsam belirleme",
         "Mikroservis tabanlı, AI destekli proje yönetim platformu hedefi tanımlandı",
         "Tamamlandı"],
        ["Docker altyapısı",
         "PostgreSQL, RabbitMQ, Redis, MailHog, Seq ve Ollama servisleri kuruldu",
         "Tamamlandı"],
        ["IdentityService",
         "Kayıt, giriş, JWT, e-posta doğrulama, AdminUserSeeder eklendi",
         "Tamamlandı"],
        ["Organizasyon ve davet",
         "Org modeli, LastActiveOrg claim'i, davet maili ve katılım akışı eklendi",
         "Tamamlandı"],
        ["Project / Sprint / Issue",
         "Üç servis CRUD, denormalize OrganizationId, Outbox + RabbitMQ event'leri",
         "Tamamlandı"],
        ["NotificationService",
         "Event tüketimi, SignalR hub, frontend bildirim merkezi",
         "Tamamlandı"],
        ["StorageService",
         "Dosya/ek yükleme, indirme, issue ile ilişkilendirme",
         "Tamamlandı"],
        ["AiService temel uçlar",
         "chat, generate-plan, enrich-issue, sprint-risk, suggest-balance, retrospective",
         "Tamamlandı"],
        ["Agent (tool-calling)",
         "AgentLoop, beş araç (CreateIssue/Sprint, AddIssueToSprint, GetActiveSprint, GetProjectIssues), audit",
         "Tamamlandı"],
        ["AI Asistan UI",
         "Scaffold + Agent moduyla iki başlıklı tek sayfa",
         "Tamamlandı"],
        ["İnce ayar altyapısı",
         "Veri toplama (Groq/Gemini), eğitim hazırlık, eval runner, deploy Modelfile",
         "Tamamlandı (kod)"],
        ["Sentetik veri üretimi",
         "Üç özellikte yaklaşık dokuz yüz örnek hedefi, ara raporda yaklaşık dört yüz örnek",
         "Devam ediyor"],
        ["Fine-tune eğitimi",
         "Colab üzerinde LoRA eğitimi, sonrasında bp-agent modeline geçiş",
         "Veri tamamlanınca"],
        ["Genel kararlılık ve UX",
         "Daha fazla senaryoda test, hata mesajları, raporlama iyileştirmeleri",
         "Devam ediyor"],
    ]
)
add_para(doc,
    "Bu aşamaya kadar yapılan çalışmalar, projenin ana omurgasının "
    "kurulduğunu göstermektedir. Bununla birlikte özellikle ince ayarlı "
    "modelin entegre edilmesi, daha geniş kullanıcı senaryolarında test "
    "edilmesi ve raporlama çıktılarının zenginleştirilmesi hâlen "
    "geliştirme gerektiren alanlardır."
)

doc.add_heading("4.2. Güncel arayüz ve çalışma örnekleri", level=2)
add_para(doc,
    "Ara rapor aşamasında elde edilen mevcut sistem çıktıları aşağıda "
    "görsellerle sunulmuştur. Bu görseller, sistemin yalnızca teorik olarak "
    "tasarlanmadığını; arayüz, servis entegrasyonu, gerçek zamanlı "
    "bildirim ve yapay zeka katmanı ile birlikte uçtan uca "
    "çalıştırılabildiğini göstermektedir."
)
sekil_aciklamalari = [
    ("Şekil 4.1. Giriş ve kayıt ekranı",
     "Kullanıcı kayıt ve giriş ekranı IdentityService ile entegre çalışmaktadır. "
     "Form doğrulama, hata mesajları ve JWT token alma akışı bu ekranda test "
     "edilmektedir."),
    ("Şekil 4.2. Organizasyon onboarding ve davet akışı",
     "Yeni kullanıcının organizasyon oluşturma veya davet kabul etme adımı "
     "MailHog ile birlikte uçtan uca doğrulanmaktadır."),
    ("Şekil 4.3. Proje listesi ve oluşturma ekranı",
     "Kullanıcının ait olduğu organizasyondaki projeler listelenmekte; yeni "
     "proje oluşturma akışı ProjectService ile gerçekleştirilmektedir."),
    ("Şekil 4.4. Sprint planlama ve takip ekranı",
     "Aktif ve tamamlanmış sprintler, sprint-issue ilişkisi ve süre takibi "
     "SprintService üzerinde yönetilmektedir."),
    ("Şekil 4.5. Kanban görünümlü issue panosu",
     "Issue'lar Open/InProgress/Done sütunları arasında sürüklenip "
     "bırakılarak hareket ettirilmekte ve her hareket SignalR ile diğer "
     "kullanıcılara da yansıtılmaktadır."),
    ("Şekil 4.6. Issue detay paneli ve AI zenginleştirme",
     "Issue detayında AiService.enrich-issue uç noktası ile başlık ve "
     "açıklama otomatik olarak zenginleştirilebilmektedir."),
    ("Şekil 4.7. AI Asistan — yeni proje taslağı (scaffold) ekranı",
     "Kullanıcı serbest metin girer; sistem proje + sprintler + issuelar "
     "şeklinde bir taslak üretir, kullanıcı onayladığında ilgili kayıtlar "
     "veri tabanlarına yazılır."),
    ("Şekil 4.8. AI Asistan — araç çağıran asistan (agent) ekranı",
     "Mevcut projede çalışan asistan, beş aracı çağırarak issue/sprint "
     "oluşturma, ekleme ve listeleme işlemlerini doğal dil ile "
     "yapabilmektedir. Tool çağrıları ve sonuçları kullanıcıya akordeonla "
     "gösterilmektedir."),
    ("Şekil 4.9. SignalR tabanlı bildirim merkezi",
     "Issue durum değişiklikleri ve davet olayları kullanıcının bildirim "
     "merkezine gerçek zamanlı düşmektedir."),
    ("Şekil 4.10. Docker servislerinin çalışır durumu",
     "docker compose ps çıktısı, on dört civarındaki servisin sağlıklı "
     "şekilde ayakta olduğunu göstermektedir."),
    ("Şekil 4.11. Seq üzerinden toplanan log akışı",
     "Servisler tarafından üretilen yapılandırılmış loglar Seq üzerinde tek "
     "bir noktadan izlenebilmektedir."),
    ("Şekil 4.12. RabbitMQ kuyruk yönetim ekranı",
     "Outbox tabanlı yayın sonrası ilgili kuyruklarda mesaj birikimi ve "
     "tüketim hızları gözlemlenebilmektedir."),
    ("Şekil 4.13. Sentetik veri üretim ilerleme görseli",
     "Faz 1 kapsamında Groq llama-3.3-70b modeli ile üretilen sentetik "
     "örneklerin tqdm tabanlı ilerleme çıktısı."),
]
for cap, desc in sekil_aciklamalari:
    add_para(doc, "")
    add_para(doc, cap, bold=True)
    add_para(doc, desc)

# ============= 5. PLANLANAN =============
doc.add_heading("5. FİNALE KADAR PLANLANAN GELİŞTİRMELER VE EKSİKLER",
                level=1)

doc.add_heading("5.1. Bu aşamada tamamlanan temel işlevler", level=2)
add_para(doc,
    "Ara rapor tarihine kadar kullanıcı kayıt/giriş, JWT tabanlı kimlik "
    "doğrulama, organizasyon yönetimi ve davet akışı, proje/sprint/issue "
    "CRUD işlemleri, Kanban panosu, dosya eki yükleme, SignalR tabanlı "
    "gerçek zamanlı bildirimler ve yedi farklı yapay zeka uç noktası "
    "geliştirilmiş durumdadır. Ayrıca araç çağıran asistan (agent), AI "
    "Asistan sayfası altında scaffold ve agent modlarıyla birlikte "
    "kullanıcıya açılmıştır. İnce ayar hattının altyapısı (veri "
    "toplama, eğitim hazırlık, eval runner, deploy Modelfile) "
    "kurulmuş; sentetik veri üretimi devam etmektedir."
)

doc.add_heading("5.2. Henüz tamamlanmamış veya güçlendirilmesi gereken noktalar",
                level=2)
add_para(doc,
    "Buna karşılık bazı başlıklar henüz bitmiş değildir. İlk olarak Faz 1 "
    "kapsamındaki sentetik veri üretimi yaklaşık dokuz yüz hedefli olup ara "
    "rapor tarihinde yaklaşık dört yüz örneğe ulaşılmıştır; eğitime "
    "geçilebilmesi için veri üretiminin tamamlanması gerekmektedir. İkinci "
    "olarak, üretilen verilerle Colab üzerinde LoRA eğitimi henüz "
    "yapılmamıştır; eğitimin ardından elde edilecek adapter ile "
    "Ollama üzerinde bp-agent modeli oluşturulacaktır."
)
add_para(doc,
    "Üçüncü olarak, base gemma3:4b modelinin agent senaryolarındaki "
    "zayıflıkları (eksik kapanan süslü parantez, markdown fence sarması, "
    "final cevap üretmek yerine aynı tool'u tekrar çağırma) ince ayar ile "
    "düzeltilmesi beklenen davranışlardır. Toleranslı parser ile palyatif "
    "çözüm sağlanmış olmakla birlikte, fine-tune sonrası çok-adımlı "
    "senaryoların kararlılığının ölçülmesi ve raporlanması "
    "gerekmektedir. Dördüncü olarak, Redis önbellek katmanı altyapı "
    "olarak hazır durumdadır ancak henüz aktif kullanım yoktur; sıkça "
    "okunan listelerde önbellekleme ileri fazda planlanmıştır."
)
add_para(doc,
    "Ayrıca arayüz tarafında raporlama, filtreleme ve dışa aktarma gibi "
    "kullanıcı deneyimi geliştirmeleri, admin panelinde organizasyon "
    "istatistikleri sekmesi ve daha kapsamlı yetkilendirme testleri "
    "planlanan iyileştirmeler arasındadır. Bazı authorization açıklarına "
    "karşılık denormalize OrganizationId ile fast-fail kontrolü eklenmiş "
    "olmakla birlikte, kapsamlı sızma testi henüz yapılmamıştır."
)

doc.add_heading("5.3. Finale kadar planlanan somut adımlar", level=2)
add_para(doc,
    "Final rapora kadar öncelikli olarak Faz 1 sentetik veri üretiminin "
    "tamamlanması, Colab üzerinde LoRA eğitiminin gerçekleştirilmesi, "
    "ortaya çıkan adapter ile bp-agent modelinin oluşturulması ve "
    "Ollama:UseFinetuned bayrağının açılması ile üretim ortamına "
    "geçirilmesi hedeflenmektedir. Bu adımdan sonra eval runner üzerinden "
    "base ve fine-tuned modellerin yan yana ölçülmesi, ortaya çıkan "
    "kazançların docs/ai-finetune-eval-v1.md dosyasındaki şablona "
    "yansıtılması planlanmaktadır."
)
add_para(doc,
    "Bunun yanında Redis tabanlı önbelleklemenin sıkça okunan uçlarda "
    "(proje listesi, sprint listesi, issue board) devreye alınması, "
    "admin panelinin organizasyon ve sistem istatistikleri sekmeleriyle "
    "zenginleştirilmesi, raporlamada PDF/CSV dışa aktarma desteğinin "
    "eklenmesi ve kapsamlı yetki testlerinin yapılması hedeflenmektedir. "
    "Son olarak proje dokümantasyonunun zenginleştirilmesi ve final "
    "rapora daha fazla deneysel çıktı ile karşılaştırmalı yorum "
    "eklenmesi planlanmaktadır."
)

# ============= 6. SONUÇ =============
doc.add_heading("6. SONUÇ VE ÖNERİLER", level=1)
add_para(doc,
    "Bu ara rapor kapsamında geliştirilen platformun önemli bir bölümü "
    "somutlaştırılmıştır. Çalışma yalnızca bir fikir aşamasında "
    "kalmamış; arayüzü bulunan, on dört civarında servisi Docker üzerinde "
    "ayakta tutan, kullanıcı kaydı, organizasyon, davet, proje, sprint ve "
    "issue akışlarını uçtan uca çalıştırabilen ve yedi farklı yapay zeka "
    "uç noktasını sunan bir yapıya dönüştürülmüştür. Projenin en güçlü "
    "yönü, çevik proje yönetimi ile yerel çalışan büyük dil modeli "
    "tabanlı asistanı tek çatı altında birleştirmesi ve ince ayar yoluyla "
    "alana özgü davranışların kazandırılmasını planlamasıdır."
)
add_para(doc,
    "Ara rapor aşamasındaki mevcut durum umut verici olmakla birlikte, "
    "finale kadar Faz 1 veri üretiminin tamamlanması, ince ayarlı modelin "
    "üretim ortamına entegre edilmesi ve kullanıcı deneyiminin "
    "olgunlaştırılması gerekmektedir. Özellikle agent senaryolarında "
    "fine-tune sonrası davranış kazançlarının ölçülmesi ve karşılaştırmalı "
    "olarak raporlanması, projenin akademik ve uygulamalı katkısını "
    "belirgin biçimde artıracaktır."
)

# ============= 7. KAYNAKLAR =============
doc.add_heading("7. KAYNAKLAR", level=1)
kaynaklar = [
    "[1]. Scrum.org, \"What is Scrum?\", https://www.scrum.org/resources/what-is-scrum.",
    "[2]. Atlassian, \"Agile project management\", https://www.atlassian.com/agile/project-management.",
    "[3]. M. Fowler, \"Microservices\", https://martinfowler.com/articles/microservices.html.",
    "[4]. RabbitMQ, \"RabbitMQ documentation\", https://www.rabbitmq.com/documentation.html.",
    "[5]. T. Brown et al., \"Language Models are Few-Shot Learners\", NeurIPS, 2020.",
    "[6]. Meta AI, \"Llama 3 model card\", https://ai.meta.com/llama/.",
    "[7]. Ollama, \"Ollama documentation\", https://ollama.com/.",
    "[8]. E. Hu et al., \"LoRA: Low-Rank Adaptation of Large Language Models\", ICLR, 2022.",
    "[9]. OpenAI, \"Function calling and tool use\", https://platform.openai.com/docs/guides/function-calling.",
    "[10]. Microsoft, \"ASP.NET Core documentation\", https://learn.microsoft.com/en-us/aspnet/core/.",
    "[11]. Microsoft, \"SignalR documentation\", https://learn.microsoft.com/en-us/aspnet/core/signalr/.",
    "[12]. PostgreSQL Global Development Group, \"PostgreSQL documentation\", https://www.postgresql.org/docs/.",
    "[13]. Docker, \"Docker Compose documentation\", https://docs.docker.com/compose/.",
]
for k in kaynaklar:
    add_para(doc, k)

# ============= 8. EKLER =============
doc.add_heading("8. EKLER", level=1)
doc.add_heading("8.1. Proje yapısı", level=2)
add_para(doc,
    "Proje yapısı incelendiğinde sistemin .NET 8/9 ve C# tabanlı sekiz "
    "ana servis bileşeni üzerine kurulduğu görülmektedir: "
    "IdentityService, ProjectService, SprintService, IssueService, "
    "NotificationService, StorageService, AiService ve Bff/Gateway. "
    "Servisler ASP.NET Core Web API yaklaşımıyla geliştirilmiş olup "
    "her biri Clean Architecture benzeri bir katman ayrımına sahiptir "
    "(Domain, Application, Infrastructure, Api). Veritabanı erişimi "
    "Entity Framework Core ve Npgsql üzerinden yapılmakta; veri "
    "modelleri her servis için ayrı bir PostgreSQL şemasında "
    "saklanmaktadır."
)
add_para(doc,
    "Servisler arası asenkron iletişim için RabbitMQ ve Outbox deseni "
    "kullanılmıştır. NotificationService gelen olayları tüketerek "
    "SignalR üzerinden frontend'e gerçek zamanlı bildirim göndermektedir. "
    "Frontend tarafında React, TypeScript ve Vite kullanılmış; durum "
    "yönetimi için React Query, stil için Tailwind CSS tercih edilmiştir."
)
add_para(doc,
    "Yapay zeka katmanında yerel çalışan Ollama altyapısı kullanılmış; "
    "metinsel görevler için gemma3:4b modeli temel alınmıştır. "
    "tools/ai_data_collector altında Groq ve Gemini sağlayıcılarıyla "
    "sentetik veri üretim hattı, tools/ai-finetune altında eğitim "
    "hazırlık adımları, tools/ai_eval altında ise değerlendirme "
    "runner'ı bulunmaktadır. tests/AiEvalDataset altında ise altmış "
    "civarında elle hazırlanmış altın (golden) örnek tutulmakta, bu "
    "set yalnızca eval amacıyla kullanılmaktadır. Tüm bu bileşenler "
    "Docker Compose tabanlı konteyner mimarisi içinde birlikte "
    "çalışacak şekilde yapılandırılmıştır."
)
add_para(doc, "")
add_para(doc, "EK-1. Proje klasör yapısı", bold=True)

# ============= 9. ÖZGEÇMİŞ =============
doc.add_heading("9. ÖZGEÇMİŞ", level=1)
add_para(doc, "Soyadı, adı : KARYAĞDI, Metin")
add_para(doc, "Uyruğu : T.C.")
add_para(doc, "Bölümü : Bilgisayar Mühendisliği")
add_para(doc, f"Öğrenci numarası : {STUDENT_NO}")
add_para(doc,
    "Çalışma alanı : .NET tabanlı mikroservis mimarisi, çevik proje "
    "yönetim sistemleri, yerel çalışan büyük dil modelleri ve LoRA "
    "tabanlı ince ayar")

doc.save(OUTPUT)
print(f"WROTE {OUTPUT}")
